"""Offline-Tests fuer den Anonymisierungs-Agenten.

Ersetzt den LM-Studio-Client durch einen Stub, damit Format-Erkennung,
Modus-Wahl, XLSX-/EML-Verarbeitung, Regex-Vorfilter und der Zaehler
ohne erreichbares Modell geprueft werden koennen.

Aufruf:  python test_offline.py
Der echte End-to-End-Test mit LM Studio bleibt davon unberuehrt.
"""
import contextlib
import io
import re
import sys
import tempfile
from pathlib import Path

import agent


# --- Fake-LLM -------------------------------------------------------------
class _Msg:
    def __init__(self, content):
        self.content = content


class _Choice:
    def __init__(self, content):
        self.message = _Msg(content)


class _Resp:
    def __init__(self, content):
        self.choices = [_Choice(content)]


class _Completions:
    @staticmethod
    def create(model, messages, temperature):
        """Simuliert das Modell: Guard-Szenarien + Standard-Ersetzungen.

        Der 2. Versuch nutzt PROMPT_ANTWORT_STRIKT (erkennbar am Wort 'VERBOTEN'),
        sodass der Stub pro Versuch unterschiedlich antworten kann.
        """
        inhalt = messages[0]["content"]
        strikt = "VERBOTEN" in inhalt
        text = inhalt.split("Text:\n", 1)[1] if "Text:\n" in inhalt else inhalt
        zelle = text.strip()

        # --- Guard-Test-Szenarien (eindeutige Eingaben) ---
        if zelle == "Hans Beispiel":
            # 1. Versuch erfindet Token, 2. Versuch liefert gueltigen Platzhalter
            return _Resp("[NAME]" if strikt else "[BERUF]")
        if zelle == "Steuerberater":
            # beide Versuche erfinden ein Token -> Fallback
            return _Resp("[ROHGEWINN]" if strikt else "[BERUF]")
        if zelle == "Krankenkasse":
            # 2. Versuch wirft eine Exception -> Fallback
            if strikt:
                raise ConnectionError("Simulierter Verbindungsabbruch")
            return _Resp("[BERUF]")

        # --- Volltext-Guard-Szenarien (eindeutige Absaetze) ---
        if zelle == "Absatz mit Betrag":
            # 1. Versuch erfindet Token, 2. Versuch liefert gueltigen Text
            return _Resp("Betrag gezahlt" if strikt else "[BETRAG] gezahlt")
        if zelle == "Absatz Fallback":
            # beide Versuche erfinden Tokens -> Fallback auf Original
            return _Resp("[KMU]" if strikt else "[BETRAG]")
        if zelle == "Absatz Exception":
            # 2. Versuch wirft Exception -> Fallback auf Original
            if strikt:
                raise ConnectionError("Simulierter Verbindungsabbruch")
            return _Resp("[BETRAG]")

        # --- Standardverhalten (uebrige Tests) ---
        ergebnis = re.sub(r"[\w.+-]+@[\w-]+\.[\w.]+", "[EMAIL]", text)
        ergebnis = ergebnis.replace("Max Mustermann", "[NAME]")
        ergebnis = ergebnis.replace("Erika Musterfrau", "[NAME]")
        ergebnis = ergebnis.replace("Musterfirma GmbH", "[FIRMA]")
        return _Resp(ergebnis)


class _Chat:
    completions = _Completions()


class _FakeClient:
    chat = _Chat()

    class models:
        @staticmethod
        def list():
            return []


agent.client = _FakeClient()


# --- Test-Infrastruktur ---------------------------------------------------
_ergebnisse = []


def pruefe(bedingung, beschreibung):
    status = "PASS" if bedingung else "FAIL"
    _ergebnisse.append((status, beschreibung))
    print(f"  [{status}] {beschreibung}")


def lese(pfad):
    return Path(pfad).read_text(encoding="utf-8")


# --- Tests ----------------------------------------------------------------
def test_verbindung():
    print("\n# Verbindungspruefung")
    try:
        agent.pruefe_verbindung()
        pruefe(True, "pruefe_verbindung() laeuft mit erreichbarem (Stub-)Client durch")
    except Exception as e:  # noqa: BLE001
        pruefe(False, f"pruefe_verbindung() warf {e}")


def test_txt(arbeit):
    print("\n# TXT -> Volltext-Modus")
    quelle = arbeit / "input" / "kunde.txt"
    quelle.write_text("Kontakt: Max Mustermann, max@test.de", encoding="utf-8")
    agent.verarbeite_datei(quelle)
    out = arbeit / "output" / "kunde_anonymisiert.txt"
    pruefe(out.exists(), "TXT-Ausgabe als .txt gespeichert (natives Format)")
    inhalt = lese(out)
    pruefe("[NAME]" in inhalt and "[EMAIL]" in inhalt, "Name und E-Mail ersetzt")
    pruefe("Max Mustermann" not in inhalt, "Klarname nicht mehr im Output")
    pruefe((arbeit / "processed" / "kunde.txt").exists(), "Original nach processed/ verschoben")


def test_docx_volltext(arbeit):
    print("\n# DOCX (einfach) -> Volltext-Modus")
    from docx import Document
    quelle = arbeit / "input" / "brief.docx"
    doc = Document()
    doc.add_paragraph("Sehr geehrter Herr Max Mustermann,")
    doc.add_paragraph("wir melden uns bei max@test.de.")
    doc.save(quelle)
    agent.verarbeite_datei(quelle)
    out = arbeit / "output" / "brief_anonymisiert.docx"
    pruefe(out.exists(), "DOCX-Ausgabe gespeichert")
    text = "\n".join(p.text for p in Document(out).paragraphs)
    pruefe("[NAME]" in text, "Name im DOCX-Volltext ersetzt")


def test_docx_tabelle(arbeit):
    print("\n# DOCX mit Tabelle -> Tabellen-Modus")
    from docx import Document
    quelle = arbeit / "input" / "tabelle.docx"
    doc = Document()
    doc.add_paragraph("Kundenliste:")
    tab = doc.add_table(rows=2, cols=2)
    tab.cell(0, 0).text = "Name"
    tab.cell(0, 1).text = "Kontakt"
    tab.cell(1, 0).text = "Max Mustermann"
    tab.cell(1, 1).text = "max@test.de"
    doc.save(quelle)
    agent.verarbeite_datei(quelle)
    out = arbeit / "output" / "tabelle_anonymisiert.docx"
    pruefe(out.exists(), "DOCX-Tabellen-Ausgabe gespeichert")
    erg = Document(out)
    zellen_text = "\n".join(
        z.text for t in erg.tables for r in t.rows for z in r.cells
    )
    pruefe("[NAME]" in zellen_text, "Name in Tabellenzelle anonymisiert (LLM-Pfad)")
    pruefe("[EMAIL]" in zellen_text, "E-Mail in Tabellenzelle anonymisiert (Regex-Vorfilter)")
    pruefe("Max Mustermann" not in zellen_text, "Klarname nicht mehr in der Tabelle")


def test_xlsx(arbeit):
    print("\n# XLSX -> Tabellen-Modus")
    import openpyxl
    quelle = arbeit / "input" / "kunden.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Kunde"
    ws["B1"] = "Umsatz"
    ws["C1"] = "Doppelt"
    ws["A2"] = "Max Mustermann"
    ws["B2"] = 1000
    ws["C2"] = "=B2*2"
    ws["A3"] = "max@test.de"
    ws["B3"] = 2000
    ws["C3"] = "=B3*2"
    wb.save(quelle)
    agent.verarbeite_datei(quelle)
    out = arbeit / "output" / "kunden_anonymisiert.xlsx"
    pruefe(out.exists(), "XLSX-Ausgabe gespeichert")
    erg = openpyxl.load_workbook(out).active
    pruefe(erg["A2"].value == "[NAME]", "Namens-Zelle anonymisiert (LLM-Pfad)")
    pruefe(erg["A3"].value == "[EMAIL]", "E-Mail-Zelle anonymisiert (Regex-Vorfilter)")
    pruefe(erg["C2"].value == "=B2*2", "Formel C2 unveraendert erhalten")
    pruefe(erg["B2"].value == 1000, "Zahlenwert B2 unveraendert erhalten")


def test_eml(arbeit):
    print("\n# EML -> Volltext-Modus")
    quelle = arbeit / "input" / "anfrage.eml"
    quelle.write_text(
        "From: Max Mustermann <max@test.de>\n"
        "To: Erika Musterfrau <erika@test.de>\n"
        "Subject: Angebot\n"
        "\n"
        "Hallo, bitte ein Angebot fuer Musterfirma GmbH.\n",
        encoding="utf-8",
    )
    agent.verarbeite_datei(quelle)
    out = arbeit / "output" / "anfrage_anonymisiert.txt"
    pruefe(out.exists(), "EML-Ausgabe als .txt gespeichert")
    inhalt = lese(out)
    pruefe("[NAME]" in inhalt and "[EMAIL]" in inhalt, "Absender/Empfaenger anonymisiert")
    pruefe("[FIRMA]" in inhalt, "Firma im Body anonymisiert")
    pruefe("max@test.de" not in inhalt, "Klartext-Adresse nicht mehr im Output")


def test_beschaedigt(arbeit):
    print("\n# Beschaedigte Datei -> DateiBeschaedigt")
    quelle = arbeit / "input" / "kaputt.docx"
    quelle.write_text("das ist kein gueltiges docx", encoding="utf-8")
    try:
        agent.verarbeite_datei(quelle)
        pruefe(False, "verarbeite_datei() haette DateiBeschaedigt werfen muessen")
    except agent.DateiBeschaedigt:
        pruefe(True, "DateiBeschaedigt bei kaputter DOCX ausgeloest")
    except Exception as e:  # noqa: BLE001
        pruefe(False, f"falsche Ausnahme: {type(e).__name__}: {e}")


def test_guard_retry():
    print("\n# Guard: 1. Versuch verworfen, 2. Versuch gueltig")
    stat = agent.Statistik()
    vorher = agent.FALLBACK_ZAEHLER
    ergebnis = agent._anonymisiere_zelle("Hans Beispiel", stat)
    pruefe(ergebnis == "[NAME]", "2. Versuch liefert gueltigen Platzhalter -> uebernommen")
    pruefe(agent.FALLBACK_ZAEHLER == vorher, "kein Fallback gezaehlt (Retry erfolgreich)")


def test_guard_fallback():
    print("\n# Guard: beide Versuche verworfen -> Fallback + Log")
    stat = agent.Statistik()
    vorher = agent.FALLBACK_ZAEHLER
    ergebnis = agent._anonymisiere_zelle("Steuerberater", stat)
    pruefe(ergebnis == "Steuerberater", "Zelle auf Originalwert zurueckgesetzt")
    pruefe(agent.FALLBACK_ZAEHLER == vorher + 1, "Fallback-Zaehler erhoeht")
    log = lese(agent.FALLBACK_LOG)
    pruefe("Steuerberater" in log, "Eintrag in anonymisierung_fallback.txt")
    kein_fehler_eintrag = (not agent.FEHLER_LOG.exists()) or ("Steuerberater" not in lese(agent.FEHLER_LOG))
    pruefe(kein_fehler_eintrag, "kein Eintrag im technischen Fehler-Log")


def test_guard_retry_exception():
    print("\n# Guard: 2. Versuch wirft Exception -> Fallback")
    stat = agent.Statistik()
    vorher = agent.FALLBACK_ZAEHLER
    ergebnis = agent._anonymisiere_zelle("Krankenkasse", stat)
    pruefe(ergebnis == "Krankenkasse", "Zelle auf Originalwert zurueckgesetzt")
    pruefe(agent.FALLBACK_ZAEHLER == vorher + 1, "Fallback-Zaehler erhoeht")
    log = lese(agent.FALLBACK_LOG)
    pruefe("ConnectionError" in log, "Exception-Typ im Fallback-Log vermerkt")


def test_volltext_guard_retry():
    print("\n# Volltext-Guard: 1. Versuch verworfen, 2. Versuch gueltig")
    stat = agent.Statistik()
    vorher = agent.FALLBACK_ABSATZ_ZAEHLER
    ergebnis = agent._anonymisiere_absatz("Absatz mit Betrag", stat)
    pruefe(ergebnis == "Betrag gezahlt", "2. Versuch (gueltig) wird verwendet")
    pruefe(agent.FALLBACK_ABSATZ_ZAEHLER == vorher, "kein Absatz-Fallback (Retry erfolgreich)")


def test_volltext_guard_fallback():
    print("\n# Volltext-Guard: beide Versuche verworfen -> Original unveraendert")
    stat = agent.Statistik()
    vorher = agent.FALLBACK_ABSATZ_ZAEHLER
    ergebnis = agent._anonymisiere_absatz("Absatz Fallback", stat)
    pruefe(ergebnis == "Absatz Fallback", "Original-Absatz unveraendert zurueckgegeben")
    pruefe(agent.FALLBACK_ABSATZ_ZAEHLER == vorher + 1, "Absatz-Fallback-Zaehler erhoeht")
    log = lese(agent.FALLBACK_LOG)
    pruefe("Absatz Fallback" in log, "Eintrag in anonymisierung_fallback.txt")
    pruefe("unveraendert belassen" in log, "Hinweis 'Volltext-Absatz unveraendert belassen' geloggt")


def test_volltext_guard_retry_exception():
    print("\n# Volltext-Guard: 2. Versuch wirft Exception -> Original unveraendert")
    stat = agent.Statistik()
    vorher = agent.FALLBACK_ABSATZ_ZAEHLER
    ergebnis = agent._anonymisiere_absatz("Absatz Exception", stat)
    pruefe(ergebnis == "Absatz Exception", "Original-Absatz unveraendert zurueckgegeben")
    pruefe(agent.FALLBACK_ABSATZ_ZAEHLER == vorher + 1, "Absatz-Fallback-Zaehler erhoeht")
    log = lese(agent.FALLBACK_LOG)
    pruefe("ConnectionError" in log, "Exception-Typ im Fallback-Log vermerkt")


def test_cloud_modell_erkennung():
    print("\n# Cloud-Schutz: Cloud-Modelle werden ausgeblendet")
    original = agent._verfuegbare_modelle
    agent._verfuegbare_modelle = lambda: ["google/gemma-3-4b", "gpt-4o", "claude-3-5-sonnet"]
    try:
        puffer = io.StringIO()
        with contextlib.redirect_stdout(puffer):
            gewaehlt = agent.waehle_modell()  # nicht-interaktiv -> Default nach Filterung
        ausgabe = puffer.getvalue()
    finally:
        agent._verfuegbare_modelle = original
    pruefe(gewaehlt == "google/gemma-3-4b", "nur lokales Modell wird gewaehlt")
    pruefe("ausgeblendet" in ausgabe, "Hinweis-Block fuer ausgeblendete Modelle erscheint")
    pruefe("gpt-4o" in ausgabe and "claude-3-5-sonnet" in ausgabe, "Cloud-Modelle im Hinweis genannt")


def test_keine_lokalen_modelle():
    print("\n# Cloud-Schutz: keine lokalen Modelle -> harter Stopp")
    original = agent._verfuegbare_modelle
    agent._verfuegbare_modelle = lambda: ["gpt-4o", "claude-3-5-sonnet"]
    try:
        with contextlib.redirect_stdout(io.StringIO()):
            agent.waehle_modell()
        pruefe(False, "haette mit sys.exit(1) abbrechen muessen")
    except SystemExit as e:
        pruefe(e.code == 1, "harter Stopp mit Exit-Code 1")
    finally:
        agent._verfuegbare_modelle = original


def test_config_model_ist_cloud():
    print("\n# Cloud-Schutz: hartkodiertes Cloud-MODEL -> harter Stopp")
    orig_flag = agent.MODELL_BEIM_START_WAEHLEN
    orig_model = agent.MODEL
    agent.MODELL_BEIM_START_WAEHLEN = False
    agent.MODEL = "gpt-4o"
    try:
        with contextlib.redirect_stdout(io.StringIO()):
            agent.waehle_modell()
        pruefe(False, "haette mit sys.exit(1) abbrechen muessen")
    except SystemExit as e:
        pruefe(e.code == 1, "harter Stopp mit Exit-Code 1")
    finally:
        agent.MODELL_BEIM_START_WAEHLEN = orig_flag
        agent.MODEL = orig_model


def test_regex_vorfilter():
    print("\n# Regex-Vorfilter")
    neu, treffer = agent._regex_vorfilter("Mail a@b.de Tel +49 170 1234567")
    pruefe("[EMAIL]" in neu and "[TELEFON]" in neu, "E-Mail und Telefon per Regex erkannt")
    pruefe(treffer == 2, f"genau 2 Regex-Treffer (war {treffer})")
    neu2, treffer2 = agent._regex_vorfilter("Jahr 2019, Wert 42")
    pruefe(treffer2 == 0, "kurze Zahlen (Jahr/Wert) nicht als Telefon erkannt")


def test_lock_erkennung(arbeit):
    print("\n# Lock-Erkennung")
    datei = arbeit / "input" / "offen.docx"
    datei.write_text("inhalt", encoding="utf-8")
    pruefe(not agent.ist_gesperrt(datei), "Datei ohne Lockdatei gilt als frei")
    lock = arbeit / "input" / "~$offen.docx"
    lock.write_text("", encoding="utf-8")
    pruefe(agent.ist_gesperrt(datei), "Datei mit '~$'-Lockdatei gilt als gesperrt")
    lock.unlink()


def main():
    arbeit = Path(tempfile.mkdtemp(prefix="anon_test_"))
    for unterordner in ("input", "output", "processed", "fehler"):
        (arbeit / unterordner).mkdir()

    # Agent auf das temporaere Arbeitsverzeichnis umlenken
    agent.OUTPUT_DIR = arbeit / "output"
    agent.PROCESSED_DIR = arbeit / "processed"
    agent.FEHLER_DIR = arbeit / "fehler"
    agent.FEHLER_LOG = arbeit / "anonymisierung_fehler.txt"
    agent.FALLBACK_LOG = arbeit / "anonymisierung_fallback.txt"

    print(f"Arbeitsverzeichnis: {arbeit}")

    test_verbindung()
    test_regex_vorfilter()
    test_lock_erkennung(arbeit)
    test_txt(arbeit)
    test_docx_volltext(arbeit)
    test_docx_tabelle(arbeit)
    test_xlsx(arbeit)
    test_eml(arbeit)
    test_beschaedigt(arbeit)
    test_guard_retry()
    test_guard_fallback()
    test_guard_retry_exception()
    test_volltext_guard_retry()
    test_volltext_guard_fallback()
    test_volltext_guard_retry_exception()
    test_cloud_modell_erkennung()
    test_keine_lokalen_modelle()
    test_config_model_ist_cloud()

    bestanden = sum(1 for s, _ in _ergebnisse if s == "PASS")
    gesamt = len(_ergebnisse)
    print(f"\n{'=' * 50}")
    print(f"Ergebnis: {bestanden}/{gesamt} Pruefungen bestanden")
    print(f"Session-Zaehler gesamt: {agent.SESSION_STAT.summe()} Ersetzungen")
    print("=" * 50)
    sys.exit(0 if bestanden == gesamt else 1)


if __name__ == "__main__":
    main()
