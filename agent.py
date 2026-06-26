"""Anonymisierungs-Agent.

Ueberwacht input_docs/, anonymisiert Dokumente (rot -> gruen) und
legt die Ergebnisse in output_docs/ ab.

Zwei automatisch erkannte Verarbeitungsmodi:
  - VOLLTEXT : TXT, PDF, DOCX, E-Mails        -> gesamter Inhalt
  - TABELLE  : XLSX, Tabellen in DOCX         -> nur Zellen mit Personenbezug

Konfiguration: siehe config.py
"""
import email
import re
import shutil
import sys
import time
from email import policy
from pathlib import Path

from openai import OpenAI, APIConnectionError, APITimeoutError, BadRequestError
import fitz  # pymupdf
from docx import Document
import openpyxl

from config import (
    LM_STUDIO_URL, MODEL, API_KEY, TEMPERATURE,
    INPUT_DIR, OUTPUT_DIR, PROCESSED_DIR, FEHLER_DIR, FEHLER_LOG, FALLBACK_LOG,
    POLL_SEKUNDEN, RETRY_SEKUNDEN, MODELL_BEIM_START_WAEHLEN, CLOUD_MODEL_MUSTER,
    PLATZHALTER, UNTERSTUETZTE_FORMATE,
    PROMPT_VOLLTEXT, PROMPT_VOLLTEXT_STRIKT, PROMPT_ANTWORT, PROMPT_ANTWORT_STRIKT,
)

client = OpenAI(base_url=LM_STUDIO_URL, api_key=API_KEY)

# Regex-Vorfilter fuer den Tabellen-Modus: offensichtliche Muster werden
# ohne LLM-Aufruf ersetzt (schneller bei grossen Tabellen).
RE_EMAIL = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
RE_TEL_KANDIDAT = re.compile(r"\+?\d[\d\s/().\-]{5,}\d")


# --- Eigene Fehlerklassen -------------------------------------------------
class AgentFehler(Exception):
    """Basisklasse fuer erwartbare Agent-Fehler."""


class LLMNichtErreichbar(AgentFehler):
    """LM Studio / lokales Modell antwortet nicht."""


class DateiBeschaedigt(AgentFehler):
    """Datei kann nicht geoeffnet/geparst werden."""


# Prompts liegen zentral in config.py (PROMPT_VOLLTEXT, PROMPT_ANTWORT,
# PROMPT_ANTWORT_STRIKT) und werden oben importiert.

# Erlaubte Platzhalter-Tokens (fuer den Halluzinations-Guard im Tabellen-Modus).
ERLAUBTE_PLATZHALTER = set(PLATZHALTER)
RE_TOKEN = re.compile(r"\[[A-ZÄÖÜ_]+\]")

# Session-Zaehler fuer die Abschluss-Zusammenfassung.
FALLBACK_ZAEHLER = 0          # Zellen, die nach 2 Versuchen auf den Vorfilter zurueckgesetzt wurden
FALLBACK_ABSATZ_ZAEHLER = 0   # Volltext-Absaetze, die nach 2 Versuchen unveraendert blieben
FEHLER_ZAEHLER = 0            # echte technische Fehler (Datei gesperrt, Format defekt, ...)


# --- Statistik ------------------------------------------------------------
class Statistik:
    """Zaehlt ersetzte Platzhalter (Naeherungswert: Anzahl Ersetzungen)."""

    def __init__(self):
        self.zaehler = {ph: 0 for ph in PLATZHALTER}

    def zaehle(self, text):
        """Addiert die im Text vorkommenden Platzhalter."""
        for ph in self.zaehler:
            self.zaehler[ph] += text.count(ph)

    def merge(self, andere):
        for ph in self.zaehler:
            self.zaehler[ph] += andere.zaehler[ph]

    def summe(self):
        return sum(self.zaehler.values())

    def bericht(self):
        zeilen = [
            f"  {PLATZHALTER[ph]}: {anzahl}"
            for ph, anzahl in self.zaehler.items() if anzahl
        ]
        return "\n".join(zeilen) if zeilen else "  (keine personenbezogenen Daten ersetzt)"


SESSION_STAT = Statistik()


# --- LLM ------------------------------------------------------------------
def _verbindungs_hinweis():
    """Klartext-Anleitung, wenn LM Studio nicht erreichbar ist."""
    return (
        "\n" + "!" * 60 + "\n"
        f"FEHLER: LM Studio (lokales Modell) ist nicht erreichbar.\n"
        f"Adresse: {LM_STUDIO_URL}\n\n"
        "Loesung:\n"
        "  1. LM Studio auf dem Rechner oeffnen\n"
        f"  2. Modell '{MODEL}' laden\n"
        "  3. Reiter 'Local Server' -> 'Start Server'\n"
        f"  4. Pruefen, dass Adresse und Port stimmen: {LM_STUDIO_URL}\n"
        "  5. Falls der Agent auf einem anderen Geraet laeuft: Netzwerk/Firewall pruefen\n"
        + "!" * 60
    )


def pruefe_verbindung():
    """Prueft beim Start, ob LM Studio antwortet. Wirft LLMNichtErreichbar."""
    try:
        client.models.list()
    except Exception as e:  # noqa: BLE001 - jede Stoerung gilt als 'nicht erreichbar'
        raise LLMNichtErreichbar(_verbindungs_hinweis()) from e


def _modell_hinweis():
    """Klartext-Anleitung, wenn das Modell nicht in den Speicher passt."""
    return (
        "\n" + "!" * 60 + "\n"
        f"FEHLER: Modell '{MODEL}' konnte nicht geladen werden.\n"
        "Wahrscheinliche Ursache: zu wenig Arbeitsspeicher (RAM) fuer dieses Modell.\n\n"
        "Loesung:\n"
        "  1. Speicherhungrige Programme schliessen (Browser-Tabs, andere KI-Tools)\n"
        "  2. ODER ein kleineres Modell waehlen (z.B. ein 1B-/3B-Modell)\n"
        "  3. In LM Studio das kleinere Modell laden und den Agenten neu starten\n"
        + "!" * 60
    )


def anonymisiere_text(text, prompt=PROMPT_VOLLTEXT, stat=None, temperature=None):
    """Sendet Text an das lokale Modell und gibt den anonymisierten Text zurueck.

    Zaehlt die ersetzten Platzhalter in stat mit (falls uebergeben).
    temperature=None -> Standardwert aus config.TEMPERATURE.
    """
    if not text or not text.strip():
        return text
    try:
        antwort = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt + text}],
            temperature=TEMPERATURE if temperature is None else temperature,
        )
    except (APIConnectionError, APITimeoutError) as e:
        raise LLMNichtErreichbar(_verbindungs_hinweis()) from e
    except BadRequestError as e:
        # Typischer Fall: Modell zu gross fuer den verfuegbaren Speicher (HTTP 400).
        raise LLMNichtErreichbar(_modell_hinweis()) from e
    ergebnis = antwort.choices[0].message.content or ""
    if stat is not None:
        stat.zaehle(ergebnis)
    return ergebnis


def _nur_erlaubte_platzhalter(text):
    """True, wenn der Text nur erlaubte Platzhalter enthaelt und kein 'UNVERAENDERT'.

    Faengt halluzinierte Tokens des Modells ab (z.B. [BERUF], [TEXT], [ROHGEWINN])
    sowie das versehentlich uebernommene Anweisungswort UNVERAENDERT.
    """
    if "UNVERAENDERT" in text:
        return False
    return all(token in ERLAUBTE_PLATZHALTER for token in RE_TOKEN.findall(text))


# --- Regex-Vorfilter (Tabellen-Modus) ------------------------------------
def _regex_vorfilter(text):
    """Ersetzt offensichtliche E-Mails und Telefonnummern per Regex.

    Gibt (neuer_text, anzahl_treffer) zurueck.
    """
    treffer = 0

    def _email_sub(_m):
        nonlocal treffer
        treffer += 1
        return "[EMAIL]"

    neu = RE_EMAIL.sub(_email_sub, text)

    def _tel_sub(m):
        nonlocal treffer
        ziffern = sum(c.isdigit() for c in m.group())
        if 7 <= ziffern <= 15:          # plausible Telefonnummer
            treffer += 1
            return "[TELEFON]"
        return m.group()

    neu = RE_TEL_KANDIDAT.sub(_tel_sub, neu)
    return neu, treffer


def _ohne_platzhalter(text):
    """Entfernt alle Platzhalter-Token aus dem Text (fuer die Rest-Pruefung)."""
    rest = text
    for ph in PLATZHALTER:
        rest = rest.replace(ph, "")
    return rest


def _anonymisiere_zelle(text, stat):
    """Anonymisiert eine einzelne Tabellen-/XLSX-Zelle.

    Strategie: Regex-Vorfilter fuer E-Mail/Telefon, danach LLM nur, wenn noch
    mehrdeutiger Text (moeglicher Name/Firma/Adresse) uebrig ist.

    Halluzinations-Guard mit Retry:
      1. Versuch  : PROMPT_ANTWORT (Standard-Temperatur)
      2. Versuch  : PROMPT_ANTWORT_STRIKT, temperature=0 (nur falls 1. verworfen/Fehler)
      Fallback    : Rueckfall auf den Regex-Vorfilter; Eintrag in anonymisierung_fallback.txt
    Jeder Aufruf ist in try/except gewickelt; es gibt keinen 3. Versuch.
    """
    global FALLBACK_ZAEHLER
    vorgefiltert, treffer = _regex_vorfilter(text)
    rest = _ohne_platzhalter(vorgefiltert)

    if not re.search(r"[A-Za-zÄÖÜäöüß]", rest):
        # Nur Zahlen/Platzhalter uebrig -> kein LLM noetig
        if treffer and stat is not None:
            stat.zaehle(vorgefiltert)
        return vorgefiltert

    # --- 1. Versuch -------------------------------------------------------
    antwort1 = None
    try:
        antwort1 = anonymisiere_text(vorgefiltert, PROMPT_ANTWORT)
        if _nur_erlaubte_platzhalter(antwort1):
            if stat is not None:
                stat.zaehle(antwort1)
            return antwort1
    except LLMNichtErreichbar:
        raise  # Modell weg / zu gross -> nicht als Zell-Fallback behandeln
    except Exception:  # noqa: BLE001 - 1. Versuch gescheitert -> 2. Versuch
        antwort1 = None

    # --- 2. Versuch (strikter Prompt, temperature=0) ----------------------
    try:
        antwort2 = anonymisiere_text(vorgefiltert, PROMPT_ANTWORT_STRIKT, temperature=0)
        if _nur_erlaubte_platzhalter(antwort2):
            if stat is not None:
                stat.zaehle(antwort2)
            return antwort2
        # Guard verwirft auch 2. Versuch -> Fallback
        log_fallback(text, antwort1, antwort2, None)
    except LLMNichtErreichbar:
        raise
    except Exception as e:  # noqa: BLE001 - 2. Versuch mit Exception -> Fallback
        log_fallback(text, antwort1, None, e)

    FALLBACK_ZAEHLER += 1
    if stat is not None and treffer:
        stat.zaehle(vorgefiltert)
    return vorgefiltert


def _anonymisiere_absatz(absatz_text, stat=None):
    """Anonymisiert einen einzelnen VOLLTEXT-Absatz mit Guard + Retry.

    Logik analog _anonymisiere_zelle, aber OHNE Regex-Vorfilter:
      1. Versuch : PROMPT_VOLLTEXT
      2. Versuch : PROMPT_VOLLTEXT_STRIKT, temperature=0 (nur falls 1. verworfen/Fehler)
      Fallback   : ORIGINAL-Absatz unveraendert; Eintrag in anonymisierung_fallback.txt
    Kein Vorfilter -> bei Fallback bleibt der Klartext stehen (bewusst; Logging + Zaehler
    + manuelle Pruefung fangen das ab). Kein 3. Versuch, keine Endlosschleife.
    """
    global FALLBACK_ABSATZ_ZAEHLER
    if not absatz_text or not absatz_text.strip():
        return absatz_text

    # --- 1. Versuch -------------------------------------------------------
    antwort1 = None
    try:
        antwort1 = anonymisiere_text(absatz_text, PROMPT_VOLLTEXT)
        if _nur_erlaubte_platzhalter(antwort1):
            if stat is not None:
                stat.zaehle(antwort1)
            return antwort1
    except LLMNichtErreichbar:
        raise  # Modell weg / zu gross -> nicht als Absatz-Fallback behandeln
    except Exception:  # noqa: BLE001 - 1. Versuch gescheitert -> 2. Versuch
        antwort1 = None

    # --- 2. Versuch (strikter Prompt, temperature=0) ----------------------
    try:
        antwort2 = anonymisiere_text(absatz_text, PROMPT_VOLLTEXT_STRIKT, temperature=0)
        if _nur_erlaubte_platzhalter(antwort2):
            if stat is not None:
                stat.zaehle(antwort2)
            return antwort2
        # Guard verwirft auch 2. Versuch -> Fallback auf Original
        log_fallback(absatz_text, antwort1, antwort2, None,
                     hinweis="Volltext-Absatz unveraendert belassen")
    except LLMNichtErreichbar:
        raise
    except Exception as e:  # noqa: BLE001 - 2. Versuch mit Exception -> Fallback
        log_fallback(absatz_text, antwort1, None, e,
                     hinweis="Volltext-Absatz unveraendert belassen")

    FALLBACK_ABSATZ_ZAEHLER += 1
    return absatz_text


def _anonymisiere_volltext(text, stat=None):
    """Splittet Text an doppelten Zeilenumbruechen und anonymisiert jeden Absatz.

    Leerzeilen-Trenner und Reihenfolge bleiben exakt erhalten; nur die inhaltlichen
    Absaetze laufen durch _anonymisiere_absatz (Guard + Retry + Fallback).
    """
    teile = re.split(r"(\n\s*\n)", text)
    return "".join(
        teil if not teil.strip() else _anonymisiere_absatz(teil, stat)
        for teil in teile
    )


# --- DOCX-Hilfen ----------------------------------------------------------
def _setze_absatz_text(para, neuer_text):
    """Schreibt neuer_text in einen Absatz und erhaelt den Absatz-Stil."""
    if not para.runs:
        para.add_run(neuer_text)
        return
    para.runs[0].text = neuer_text
    for run in para.runs[1:]:
        run.text = ""


def _setze_zellen_text(zelle, neuer_text):
    """Schreibt neuer_text in eine Tabellenzelle (erhaelt Zellen-Formatierung)."""
    _setze_absatz_text(zelle.paragraphs[0], neuer_text)
    for para in zelle.paragraphs[1:]:
        for run in para.runs:
            run.text = ""


# --- Modus: Tabelle (DOCX-Tabellen) --------------------------------------
def _docx_tabellen_anonymisieren(doc, stat):
    """Anonymisiert Zellen mit Personenbezug in allen DOCX-Tabellen."""
    for t_idx, tabelle in enumerate(doc.tables, 1):
        for zeile in tabelle.rows:
            for zelle in zeile.cells:
                text = zelle.text.strip()
                if not text:
                    continue
                neu = _anonymisiere_zelle(text, stat)
                if neu != text:
                    _setze_zellen_text(zelle, neu)
        print(f"    Tabelle {t_idx}: {len(tabelle.rows)} Zeilen geprueft")


# --- Modus: Volltext (DOCX-Absaetze, formatierungserhaltend) -------------
def _docx_absaetze_volltext(doc, stat):
    """Anonymisiert jeden Absatz einzeln (Guard + Retry) und erhaelt die Formatierung."""
    absaetze = [p for p in doc.paragraphs if p.text.strip()]
    for i, para in enumerate(absaetze, 1):
        print(f"    Absatz {i}/{len(absaetze)} ...")
        neu = _anonymisiere_absatz(para.text, stat)
        _setze_absatz_text(para, neu)


# --- Format-Handler -------------------------------------------------------
def _verarbeite_txt(pfad, output, stat):
    print("  Modus: VOLLTEXT (TXT)")
    try:
        text = pfad.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = pfad.read_text(encoding="latin-1")
        except Exception as e:  # noqa: BLE001
            raise DateiBeschaedigt(f"TXT nicht lesbar: {e}") from e
    print(f"  {len(text)} Zeichen")
    output.write_text(_anonymisiere_volltext(text, stat), encoding="utf-8")


def _verarbeite_pdf(pfad, output, stat):
    print("  Modus: VOLLTEXT (PDF)")
    try:
        doc = fitz.open(pfad)
    except Exception as e:  # noqa: BLE001
        raise DateiBeschaedigt(f"PDF nicht lesbar: {e}") from e
    teile = []
    anzahl = doc.page_count
    for i, seite in enumerate(doc, 1):
        print(f"  Seite {i}/{anzahl} ...")
        text = seite.get_text()
        teile.append(_anonymisiere_volltext(text, stat) if text.strip() else "")
    doc.close()
    output.write_text("\n\n".join(teile), encoding="utf-8")


def _verarbeite_docx(pfad, output, stat):
    try:
        doc = Document(pfad)
    except Exception as e:  # noqa: BLE001
        raise DateiBeschaedigt(f"DOCX nicht lesbar: {e}") from e

    hat_tabellen = len(doc.tables) > 0

    if hat_tabellen:
        print(f"  Modus: TABELLE ({len(doc.tables)} Tabelle(n)) + Volltext-Absaetze")
        _docx_tabellen_anonymisieren(doc, stat)
        _docx_absaetze_volltext(doc, stat)
    else:
        print("  Modus: VOLLTEXT (DOCX)")
        _docx_absaetze_volltext(doc, stat)
    doc.save(output)


def _verarbeite_xlsx(pfad, output, stat):
    print("  Modus: TABELLE (XLSX)")
    try:
        wb = openpyxl.load_workbook(pfad)  # data_only=False -> Formeln bleiben erhalten
    except Exception as e:  # noqa: BLE001
        raise DateiBeschaedigt(f"XLSX nicht lesbar: {e}") from e
    for ws in wb.worksheets:
        print(f"  Blatt '{ws.title}': {ws.max_row} Zeilen x {ws.max_column} Spalten")
        geprueft = 0
        for zeile in ws.iter_rows():
            for zelle in zeile:
                wert = zelle.value
                if wert is None or zelle.data_type == "f":
                    continue  # leer oder Formel -> unangetastet lassen
                if not isinstance(wert, str) or not wert.strip():
                    continue  # Zahl/Datum/leer -> kein anonymisierbarer Text
                neu = _anonymisiere_zelle(wert.strip(), stat)
                if neu != wert:
                    zelle.value = neu
                geprueft += 1
        print(f"    {geprueft} Text-Zellen geprueft")
    wb.save(output)


def _lese_eml(pfad):
    """Liest Header (Von/An/Cc/Betreff) und Klartext-Body aus einer .eml-Datei."""
    try:
        with open(pfad, "rb") as f:
            msg = email.message_from_binary_file(f, policy=policy.default)
    except Exception as e:  # noqa: BLE001
        raise DateiBeschaedigt(f"EML nicht lesbar: {e}") from e
    kopf = {
        "Von": str(msg.get("From", "")),
        "An": str(msg.get("To", "")),
        "Cc": str(msg.get("Cc", "")),
        "Betreff": str(msg.get("Subject", "")),
    }
    koerper = ""
    try:
        if msg.is_multipart():
            for teil in msg.walk():
                if teil.get_content_type() == "text/plain":
                    koerper += teil.get_content()
        else:
            koerper = msg.get_content()
    except Exception:  # noqa: BLE001 - exotische Encodings tolerieren
        koerper = msg.get_payload(decode=False) or ""
    return kopf, koerper


def _lese_msg(pfad):
    """Liest Header und Body aus einer Outlook .msg-Datei (Lib: extract-msg)."""
    try:
        import extract_msg
    except ImportError as e:
        raise AgentFehler(
            "Bibliothek 'extract-msg' fehlt. Installieren: pip install extract-msg"
        ) from e
    try:
        m = extract_msg.Message(str(pfad))
        kopf = {
            "Von": m.sender or "",
            "An": m.to or "",
            "Cc": m.cc or "",
            "Betreff": m.subject or "",
        }
        koerper = m.body or ""
        m.close()
    except Exception as e:  # noqa: BLE001
        raise DateiBeschaedigt(f"MSG nicht lesbar: {e}") from e
    return kopf, koerper


def _verarbeite_email(pfad, output, stat, suffix):
    print(f"  Modus: VOLLTEXT (E-Mail {suffix})")
    kopf, koerper = _lese_eml(pfad) if suffix == ".eml" else _lese_msg(pfad)
    print(f"  Header + {len(koerper)} Zeichen Body")
    anon_kopf = {
        feld: (anonymisiere_text(wert, PROMPT_ANTWORT, stat) if wert.strip() else "")
        for feld, wert in kopf.items()
    }
    anon_body = _anonymisiere_volltext(koerper, stat)
    zeilen = [f"{feld}: {anon_kopf[feld]}" for feld in kopf]
    zeilen += ["", anon_body]
    output.write_text("\n".join(zeilen), encoding="utf-8")


# Endung -> (Handler, Ausgabe-Endung)
_HANDLER = {
    ".txt": (_verarbeite_txt, ".txt"),
    ".pdf": (_verarbeite_pdf, ".txt"),
    ".docx": (_verarbeite_docx, ".docx"),
    ".xlsx": (_verarbeite_xlsx, ".xlsx"),
}


def verarbeite_datei(pfad):
    """Verarbeitet eine einzelne Datei je nach Format und Modus."""
    pfad = Path(pfad)
    suffix = pfad.suffix.lower()
    print(f"\n{'=' * 50}\nVerarbeite: {pfad.name}")

    datei_stat = Statistik()

    if suffix in _HANDLER:
        handler, out_suffix = _HANDLER[suffix]
        output = OUTPUT_DIR / (pfad.stem + "_anonymisiert" + out_suffix)
        handler(pfad, output, datei_stat)
    elif suffix in (".eml", ".msg"):
        output = OUTPUT_DIR / (pfad.stem + "_anonymisiert.txt")
        _verarbeite_email(pfad, output, datei_stat, suffix)
    else:
        print(f"  Format nicht unterstuetzt: {suffix} -> uebersprungen")
        return

    print(f"Gespeichert: output_docs/{output.name}")
    print(f"Zusammenfassung fuer {pfad.name}:")
    print(datei_stat.bericht())

    SESSION_STAT.merge(datei_stat)
    shutil.move(str(pfad), str(PROCESSED_DIR / pfad.name))
    print("Original verschoben nach: processed/")


# --- Fehlerbehandlung ----------------------------------------------------
def ist_gesperrt(pfad):
    """Prueft, ob eine Datei gesperrt ist (z.B. in Word/Excel geoeffnet).

    Word/Excel legen beim Oeffnen eine Besitzer-Datei '~$name' im selben
    Ordner an - das ist auf macOS das zuverlaessige Signal.
    """
    lockdatei = pfad.parent / ("~$" + pfad.name)
    if lockdatei.exists():
        return True
    try:
        with open(pfad, "rb"):
            return False
    except (PermissionError, OSError):
        return True


def log_fehler(kontext, exception):
    """Schreibt einen ECHTEN technischen Fehler in anonymisierung_fehler.txt.

    Fuer Datei-/Format-/Verbindungsfehler ausserhalb des Tabellen-Guards.
    Halluzinations-Rollbacks gehen NICHT hierher, sondern in log_fallback().
    """
    global FEHLER_ZAEHLER
    FEHLER_ZAEHLER += 1
    zeitpunkt = time.strftime("%Y-%m-%d %H:%M:%S")
    eintrag = f"[{zeitpunkt}] {kontext}: {type(exception).__name__}: {exception}\n"
    try:
        with open(FEHLER_LOG, "a", encoding="utf-8") as f:
            f.write(eintrag)
    except Exception as e:  # noqa: BLE001
        print(f"  WARNUNG: Fehler-Log nicht schreibbar: {e}")
    print(f"  Fehler protokolliert in {FEHLER_LOG.name}")


def log_fallback(zelle, modellantwort_versuch1, modellantwort_versuch2,
                 exception_falls_vorhanden, hinweis="Zelle auf Vorfilter zurueckgesetzt"):
    """Protokolliert einen Halluzinations-Rollback in anonymisierung_fallback.txt.

    Wird gerufen, wenn beide LLM-Versuche fuer eine Tabellenzelle bzw. einen
    Volltext-Absatz verworfen wurden (erfundene Tokens) oder eine Exception warfen.
    Der Inhalt wurde zurueckgesetzt (Zelle -> Vorfilter, Absatz -> Original) und sollte
    manuell geprueft werden. 'hinweis' beschreibt die Art des Rollbacks.
    """
    zeitpunkt = time.strftime("%Y-%m-%d %H:%M:%S")
    zeilen = [
        f"[{zeitpunkt}] {hinweis} (bitte manuell pruefen)",
        f"    Inhalt    : {zelle!r}",
        f"    Versuch 1  : {modellantwort_versuch1!r}",
        f"    Versuch 2  : {modellantwort_versuch2!r}",
    ]
    if exception_falls_vorhanden is not None:
        zeilen.append(
            f"    Exception  : {type(exception_falls_vorhanden).__name__}: {exception_falls_vorhanden}"
        )
    eintrag = "\n".join(zeilen) + "\n"
    try:
        with open(FALLBACK_LOG, "a", encoding="utf-8") as f:
            f.write(eintrag)
    except Exception as e:  # noqa: BLE001
        print(f"  WARNUNG: Fallback-Log nicht schreibbar: {e}")


def _nach_fehler(pfad):
    """Verschiebt eine beschaedigte Datei in den Ordner fehler/."""
    try:
        shutil.move(str(pfad), str(FEHLER_DIR / pfad.name))
        print(f"  Beschaedigte Datei verschoben nach: fehler/{pfad.name}")
    except Exception as e:  # noqa: BLE001
        print(f"  WARNUNG: Datei konnte nicht nach fehler/ verschoben werden: {e}")


# --- Modellauswahl --------------------------------------------------------
def _verfuegbare_modelle():
    """Liest die in LM Studio verfuegbaren Modell-IDs (leer bei Fehler)."""
    try:
        return [m.id for m in client.models.list().data]
    except Exception:  # noqa: BLE001 - keine Liste -> Default verwenden
        return []


def _ist_cloud_modell(modell_id):
    """True, wenn die Modell-ID auf ein Cloud-Modell hindeutet (einfacher Substring-Match)."""
    mid = modell_id.lower()
    return any(muster in mid for muster in CLOUD_MODEL_MUSTER)


def _keine_lokalen_stopp():
    """Harter Stopp: in LM Studio sind nur Cloud-Modelle verfuegbar."""
    print("\n" + "!" * 60)
    print("Fehler: Keine lokalen Modelle in LM Studio gefunden. Dieser Agent verarbeitet")
    print("Daten ausschliesslich lokal. Bitte lade ein lokales GGUF-Modell in LM Studio")
    print("und starte neu. Anleitung: docs/Anleitung.pdf")
    print("!" * 60)
    sys.exit(1)


def _cloud_stopp_hartkodiert(modell):
    """Harter Stopp: das in config.py hartkodierte Modell ist ein Cloud-Modell."""
    print("\n" + "!" * 60)
    print(f"Fehler: Das in config.py hartkodierte Modell '{modell}' wurde als Cloud-Modell")
    print("erkannt. Dieser Agent verarbeitet Daten ausschliesslich lokal. Bitte trage ein")
    print("lokales GGUF-Modell in MODEL ein (config.py) oder aktiviere MODELL_BEIM_START_WAEHLEN.")
    print("!" * 60)
    sys.exit(1)


def waehle_modell():
    """Laesst den Nutzer beim Start ein lokales Modell aus LM Studio waehlen.

    Cloud-Modelle (OpenAI/Anthropic/Google/...) werden hart blockiert und nicht
    angeboten - der Schutz laeuft auch headless (ohne Terminal). Faellt auf
    config.MODEL zurueck, wenn die Auswahl deaktiviert ist oder keine Modellliste
    geladen werden kann.
    """
    # B.4: hartkodiertes Modell pruefen, wenn die Auswahl deaktiviert ist.
    if not MODELL_BEIM_START_WAEHLEN:
        if _ist_cloud_modell(MODEL):
            _cloud_stopp_hartkodiert(MODEL)
        return MODEL

    modelle = _verfuegbare_modelle()
    if not modelle:
        return MODEL

    # Cloud-Modelle hart blockieren (unabhaengig von TTY).
    lokal = [m for m in modelle if not _ist_cloud_modell(m)]
    geblockt = [m for m in modelle if _ist_cloud_modell(m)]
    if geblockt:
        print("\n" + "!" * 60)
        print("WARNUNG: Folgende Modelle wurden ausgeblendet, weil sie als Cloud-Modelle")
        print("erkannt wurden:")
        for m in geblockt:
            print(f"  - {m}")
        print("Dieser Agent erlaubt ausschliesslich lokale Modelle. Falls das ein Irrtum")
        print("ist, pruefe CLOUD_MODEL_MUSTER in config.py.")
        print("!" * 60)
    if not lokal:
        _keine_lokalen_stopp()

    # Default vorne einsortieren.
    default = MODEL if MODEL in lokal else lokal[0]
    if not sys.stdin.isatty():
        return default

    print("\nVerfuegbare Modelle in LM Studio:")
    for i, mid in enumerate(lokal, 1):
        markierung = "  (Default)" if mid == default else ""
        print(f"  {i}) {mid}{markierung}")

    try:
        eingabe = input(f"Wahl [Enter = {default}]: ").strip()
    except EOFError:
        return default
    if not eingabe:
        return default
    if eingabe.isdigit() and 1 <= int(eingabe) <= len(lokal):
        return lokal[int(eingabe) - 1]
    if eingabe in lokal:
        return eingabe
    print(f"  Ungueltige Eingabe - verwende Default: {default}")
    return default


# --- Hauptschleife --------------------------------------------------------
def main():
    global MODEL
    for ordner in (INPUT_DIR, OUTPUT_DIR, PROCESSED_DIR, FEHLER_DIR):
        ordner.mkdir(exist_ok=True)

    print("=" * 50)
    print("Anonymisierungs-Agent gestartet")
    print(f"Unterstuetzte Formate: {', '.join(sorted(UNTERSTUETZTE_FORMATE))}")
    print("Modi: VOLLTEXT | TABELLE (automatisch erkannt)")
    print("Zum Beenden: Strg + C")
    print("=" * 50)

    try:
        pruefe_verbindung()
    except LLMNichtErreichbar as e:
        print(str(e))
        return

    MODEL = waehle_modell()
    print(f"\nLM Studio erreichbar - Modell: {MODEL}\n")

    print(f"Ueberwache {INPUT_DIR.name}/ ...")
    verarbeitet = set()
    retry_queue = {}  # Dateiname -> fruehester naechster Versuch (Zeitstempel)

    try:
        while True:
            jetzt = time.time()
            for datei in sorted(INPUT_DIR.iterdir()):
                if not datei.is_file():
                    continue
                if datei.name.startswith(".") or datei.name.startswith("~$"):
                    continue  # versteckte Dateien / Office-Lockdateien ignorieren
                if datei.name in verarbeitet:
                    continue
                if datei.name in retry_queue and jetzt < retry_queue[datei.name]:
                    continue
                if datei.suffix.lower() not in UNTERSTUETZTE_FORMATE:
                    print(f"Uebersprungen (Format nicht unterstuetzt): {datei.name}")
                    verarbeitet.add(datei.name)
                    continue
                if ist_gesperrt(datei):
                    retry_queue[datei.name] = jetzt + RETRY_SEKUNDEN
                    print(f"Datei gesperrt (in Word/Excel geoeffnet?): {datei.name} "
                          f"- neuer Versuch in {RETRY_SEKUNDEN}s")
                    continue

                retry_queue.pop(datei.name, None)
                verarbeitet.add(datei.name)
                try:
                    verarbeite_datei(datei)
                except DateiBeschaedigt as e:
                    log_fehler(datei.name, e)
                    _nach_fehler(datei)
                except LLMNichtErreichbar as e:
                    print(str(e))
                    verarbeitet.discard(datei.name)
                    print("\nAgent pausiert. LM Studio starten, dann Agent neu ausfuehren.")
                    return
                except Exception as e:  # noqa: BLE001
                    log_fehler(datei.name, e)
                    print(f"Uebersprungen wegen Fehler: {datei.name}")
            time.sleep(POLL_SEKUNDEN)
    except KeyboardInterrupt:
        print("\n" + "=" * 50)
        print("Agent beendet - Sitzungs-Zusammenfassung:")
        print(SESSION_STAT.bericht())
        print(f"Insgesamt ersetzt: {SESSION_STAT.summe()} Eintraege")
        print(f"{FEHLER_ZAEHLER} echte Fehler (siehe {FEHLER_LOG.name})")
        print(f"{FALLBACK_ZAEHLER} Zellen mit Fallback "
              f"(siehe {FALLBACK_LOG.name} - bitte manuell pruefen)")
        print(f"{FALLBACK_ABSATZ_ZAEHLER} Absaetze mit Fallback "
              f"(siehe {FALLBACK_LOG.name} - Original unveraendert belassen, bitte manuell pruefen)")
        print("=" * 50)


if __name__ == "__main__":
    main()
